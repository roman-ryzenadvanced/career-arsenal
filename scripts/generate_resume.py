"""
Resume Architect — Phase 3-4: Generate ATS-optimized resume for Roman M.

Target role: COO / Head of Operations / VP of Customer Experience
Target company: Series A Web3 / blockchain infrastructure startup, remote

Produces both PDF (via ReportLab, following pdf/briefs/resume.md template)
and DOCX (via python-docx, ATS-safe single-column layout).

Output: /home/z/my-project/download/Roman_M_Resume_COO_Web3.pdf
        /home/z/my-project/download/Roman_M_Resume_COO_Web3.docx
"""

import os
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import cm, mm
from reportlab.lib.styles import ParagraphStyle
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_JUSTIFY
from reportlab.lib import colors
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable, KeepTogether
)
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.pdfbase.pdfmetrics import registerFontFamily

from docx import Document
from docx.shared import Pt, Cm, RGBColor, Twips
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.section import WD_SECTION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ─────────────────────────────────────────────────────────────────────────────
# RESUME CONTENT  (single source of truth for both PDF and DOCX)
# ─────────────────────────────────────────────────────────────────────────────

NAME = "Roman M."
TARGET_ROLE = "Chief Operating Officer  |  Blockchain Infrastructure & Web3 Operations"
CONTACT_LINE = "rommark@gmx.com  |  Georgia (Remote)  |  linkedin.com/in/roman-m-793b3310  |  Hebrew · English · Russian"

PROFESSIONAL_SUMMARY = (
    "Operations executive with 19+ years building and scaling customer-facing "
    "organizations across blockchain infrastructure, data center, and cloud hosting "
    "industries. Currently serving as Chief Operations Officer at DedicatedNodes, "
    "owning sales, customer operations, and talent acquisition for a blockchain "
    "infrastructure provider serving Web3 projects. Proven track record leading "
    "distributed teams across four continents, managing multi-million-dollar hosting "
    "portfolios, and architecting customer success organizations from zero to scaled "
    "operations. Deep Web3 domain expertise spanning stablecoin protocols, validator "
    "infrastructure, developer community building, and high-uptime data center operations."
)

# Skills grouped by category — keyword-matched to typical Series A Web3 COO JDs
SKILLS = [
    ("Operations & Leadership",
     "Operations Strategy, Cross-Functional Leadership, P&L Management, OKRs / KPIs, "
     "Process Optimization, SOP Development, Vendor Management, Remote-First Global Team Leadership"),
    ("Customer & Community",
     "Customer Success, Customer Experience (CX), Community Management, Developer Relations, "
     "Escalation Management, Retention Strategy, NPS / CSAT Programs, HelpDesk Operations"),
    ("Sales & GTM",
     "B2B / B2C Sales, Enterprise Account Management, Go-to-Market Strategy, "
     "Partnership Development, Revenue Operations, Google Ads, SMM / Internet Marketing"),
    ("People & Talent",
     "Talent Acquisition (HeadHunter), Team Building, Recruiting & Training, "
     "Performance Management, Remote Talent Sourcing, International Team Supervision"),
    ("Domain Expertise",
     "Blockchain Infrastructure, Web3, Crypto / Stablecoins, Data Center Operations, "
     "Colocation, Dedicated Servers, Cloud Hosting, High-Bandwidth Solutions, Fintech, Trading Systems"),
    ("Tools & Platforms",
     "Discord Community Management, HelpDesk / Ticketing Systems, Google Workspace, "
     "Microsoft Office, LinkedIn Recruiter, SMM Tools"),
]

# Work experience — reverse chronological, most recent detailed, early roles compressed
EXPERIENCE = [
    {
        "title": "Chief Operations Officer  |  Head of Sales  |  Director of Customer Operations",
        "company": "DedicatedNodes",
        "location": "Netherlands (Remote)",
        "dates": "Mar 2025 – Present",
        "bullets": [
            "Lead end-to-end operations for a blockchain infrastructure provider delivering "
            "dedicated hardware and software solutions to Web3 projects and protocols.",
            "Own sales organization, customer operations, and talent acquisition functions "
            "across a fully remote, distributed workforce spanning multiple time zones.",
            "Architect customer onboarding and escalation workflows engineered for the "
            "high-uptime demands of validator, node-hosting, and RPC endpoint clients.",
            "Drive technical and commercial partnerships with blockchain protocols and "
            "data center facilities to expand serviceable infrastructure footprint.",
        ],
    },
    {
        "title": "Ambassador (Pro Bono)",
        "company": "Z.ai",
        "location": "Georgia (Remote)",
        "dates": "Feb 2026 – Present",
        "bullets": [
            "Test and provide structured feedback on emerging AI agentic and coding "
            "models as part of the ambassador team, working directly with product and "
            "research teams to surface usability gaps and improvement opportunities.",
        ],
    },
    {
        "title": "Community Specialist (Pro Bono)",
        "company": "Trae.ai  —  Vibe Coding Community",
        "location": "Georgia (Remote)",
        "dates": "Jan 2026 – Present",
        "bullets": [
            "Manage Discord community for developers exploring AI-powered coding "
            "workflows; organize interactive discussions and collaborative events "
            "connecting learners, creators, and innovators using Trae.ai's IDE platform.",
        ],
    },
    {
        "title": "Customer Experience Manager / Supervisor / Remote Talents Hunting  "
                 "(also: Team Lead & Customer Experience Manager; Abuse Manager)",
        "company": "QuadraNet Enterprises, LLC.",
        "location": "California, USA (Remote)",
        "dates": "Feb 2019 – Mar 2025  (6 yr 2 mo)",
        "bullets": [
            "Led customer experience operations for one of Los Angeles's largest "
            "dedicated server, colocation, cloud hosting, and bandwidth providers "
            "(400+ fiber strands at the One Wilshire Meet-Me-Room).",
            "Built and supervised remote talent acquisition pipeline for IT roles — "
            "sourced technical support, abuse management, and operations staff across "
            "multiple time zones, reducing time-to-hire for technical roles.",
            "Owned Team Lead & Customer Experience Manager function: supervised "
            "international engineering and customer service teams, owned escalation "
            "management, and drove quality-of-service assurance across data center "
            "product lines.",
            "Directed Abuse Management operations: established abuse handling policies, "
            "escalation paths, and cross-functional response protocols for security "
            "and compliance incidents impacting hosted infrastructure.",
        ],
    },
    {
        "title": "Director of Customer Success",
        "company": "Kowala  (Blockchain Project — first autonomous decentralized stablecoin)",
        "location": "Remote",
        "dates": "Jan 2018 – Nov 2018",
        "bullets": [
            "Built and led customer success function for a stablecoin protocol combining "
            "Ethereum smart contracts with a Tendermint-derived Proof-of-Stake consensus "
            "(1-second finality, on-chain dynamic validator set management).",
            "Managed a team of four technical support and service personnel; owned "
            "helpdesk monitoring, escalations, and SLA performance.",
            "Served as Community Support Manager and technical assistant to the marketing "
            "director; coordinated web-conference meetings with ecosystem partners and "
            "external projects.",
        ],
    },
    {
        "title": "Director of Customer Service / Helpdesk Team Recruiting",
        "company": "Kowala  (Blockchain Project)",
        "location": "Remote",
        "dates": "Feb 2017 – Dec 2017",
        "bullets": [
            "Established customer service operations for an early-stage blockchain "
            "protocol; integrated helpdesk tooling, monitoring, and escalation workflows.",
            "Recruited and trained a technical support team of four; managed community "
            "support operations through the protocol launch phase.",
        ],
    },
    {
        "title": "Business Development  (Product Manager / Global Sales Manager / B2B Lead)",
        "company": "WIZZSOLUTIONS.COM GROUP LLC",
        "location": "Miami, FL, USA",
        "dates": "Feb 2013 – Feb 2016  (3 yr 1 mo)",
        "bullets": [
            "Led B2B and B2C sales and business development for hosting and remote IT "
            "management services, negotiating key agreements with technology partners "
            "and data center facilities worldwide.",
            "Designed commercial solutions and customer solution engineering for "
            "high-bandwidth, dedicated server, and colocation clients; built marketing "
            "strategies for global customers and strategic partners.",
            "Supervised international technical support engineering offices (India & "
            "Ukraine), managed account relationships for strategic global customers, "
            "and operated the sales department end-to-end.",
        ],
    },
    {
        "title": "Global Account & Retention Manager  /  Google Ads Administrator",
        "company": "OMC  /  Global Cloud Operator  (Headquarters Office)",
        "location": "Remote",
        "dates": "Jan 2012 – Feb 2013  (1 yr 2 mo)",
        "bullets": [
            "Managed international cloud computing accounts and retention programs; "
            "ran Google Ads promotional campaigns and supported technical escalations.",
            "Acted as Project Manager for international cloud customers — designed "
            "custom solutions, generated B2B and consumer leads, and consulted on "
            "outsourcing and bespoke product design.",
        ],
    },
    {
        "title": "International Branch Manager  |  Data Center Services Company",
        "company": "Data Center Operations Provider  (NDA)",
        "location": "Florida, USA",
        "dates": "Jan 2007 – Sep 2011  (4 yr 9 mo)",
        "bullets": [
            "Led overseas branch of a data center services company; managed onsite and "
            "remote staff, partnerships with global providers, and high-bandwidth "
            "product development.",
            "Owned sales department and contributed to marketing strategy and product "
            "design as part of internal consulting function.",
            "Branch was consolidated into the US main office as part of a strategic "
            "business optimization; oversaw migration of infrastructure and customers.",
        ],
    },
]

# Compressed earlier experience — listed without bullets to save space
EARLIER_EXPERIENCE = [
    ("Content Assistant (Pro Bono)", "Fintech Industry (NDA), Remote",
     "Jul 2016 – Aug 2016",
     "Wrote articles on manual and automated trading products; tested trading "
     "solutions under demo environments to produce objective technical feedback."),
    ("Content & Creative Consulting (Pro Bono)", "Startup Education Project, Toronto, Canada",
     "Jul 2016",
     "Authored content and advised on project structure, roadmap, and product "
     "design; assisted with private presentations used by the CEO in partner "
     "and investor meetings."),
    ("Remote IT Recruiting (Pro Bono)", "Remote Employment for IT Industry, Canada",
     "Jun 2016 – Jul 2016",
     "Sourced talent for IT companies; wrote unique recruitment-focused content "
     "for an innovative remote-first recruitment model."),
]

EDUCATION = [
    ("Online Marketing Professional  —  Internet Marketing and SMM",
     "Udemy Academy", "2018 – 2019"),
    ("Windows Server Administrator  —  Internet Technologies",
     "Federal Technological Institute", "2006 – 2007"),
    ("Technical Team Lead & Chief Of Abuse Officer  —  Information Technology",
     "Federal Duty (NDA)", "2004 – 2007"),
    ("Public Media and Production",
     "Beit Rutnerberg", "2001 – 2004"),
]

LANGUAGES = "Hebrew (Native or Bilingual)  ·  English (Full Professional)  ·  Russian (Elementary)"

# ─────────────────────────────────────────────────────────────────────────────
# PDF GENERATION  (follows pdf/briefs/resume.md ATS-safe template)
# ─────────────────────────────────────────────────────────────────────────────

def build_pdf(output_path):
    """Generate ATS-optimized PDF resume."""

    # Font registration
    pdfmetrics.registerFont(TTFont('FreeSerif',
                                   '/usr/share/fonts/truetype/freefont/FreeSerif.ttf'))
    pdfmetrics.registerFont(TTFont('FreeSerif-Bold',
                                   '/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf'))
    pdfmetrics.registerFont(TTFont('FreeSerif-Italic',
                                   '/usr/share/fonts/truetype/freefont/FreeSerifItalic.ttf'))
    pdfmetrics.registerFont(TTFont('FreeSerif-BoldItalic',
                                   '/usr/share/fonts/truetype/freefont/FreeSerifBoldItalic.ttf'))
    registerFontFamily('FreeSerif',
                       normal='FreeSerif', bold='FreeSerif-Bold',
                       italic='FreeSerif-Italic', boldItalic='FreeSerif-BoldItalic')

    # ATS-safe palette: deep navy accent (professional, blockchain/finance-coded)
    ACCENT = colors.HexColor('#1A3352')
    TEXT_MUTED = colors.HexColor('#555555')
    TEXT = colors.HexColor('#1A1A1A')

    # Styles — single-column ATS-safe, tightened to fit 2 pages
    name_style = ParagraphStyle(
        'ResumeName', fontName='FreeSerif-Bold', fontSize=22,
        leading=26, alignment=TA_CENTER, spaceAfter=2, textColor=TEXT
    )
    target_role_style = ParagraphStyle(
        'ResumeTargetRole', fontName='FreeSerif', fontSize=10.5,
        leading=13, alignment=TA_CENTER, spaceAfter=2,
        textColor=ACCENT
    )
    contact_style = ParagraphStyle(
        'ResumeContact', fontName='FreeSerif', fontSize=9,
        leading=12, alignment=TA_CENTER, textColor=TEXT_MUTED, spaceAfter=3
    )
    section_title_style = ParagraphStyle(
        'ResumeSectionTitle', fontName='FreeSerif-Bold', fontSize=12,
        leading=14, spaceBefore=7, spaceAfter=2, textColor=ACCENT
    )
    job_title_style = ParagraphStyle(
        'ResumeJobTitle', fontName='FreeSerif-Bold', fontSize=10,
        leading=12.5, spaceAfter=1, textColor=TEXT
    )
    job_meta_style = ParagraphStyle(
        'ResumeJobMeta', fontName='FreeSerif-Italic', fontSize=9,
        leading=11.5, textColor=TEXT_MUTED, spaceAfter=2
    )
    bullet_style = ParagraphStyle(
        'ResumeBullet', fontName='FreeSerif', fontSize=9.5,
        leading=12, leftIndent=12, bulletIndent=0,
        spaceBefore=0, spaceAfter=0, textColor=TEXT
    )
    body_style = ParagraphStyle(
        'ResumeBody', fontName='FreeSerif', fontSize=9.5,
        leading=12, spaceAfter=2, textColor=TEXT,
        alignment=TA_LEFT
    )
    skills_style = ParagraphStyle(
        'ResumeSkills', fontName='FreeSerif', fontSize=9.5,
        leading=12, spaceBefore=0, spaceAfter=0, textColor=TEXT,
        leftIndent=0
    )
    earlier_style = ParagraphStyle(
        'ResumeEarlier', fontName='FreeSerif', fontSize=9,
        leading=11.5, leftIndent=12, spaceBefore=0, spaceAfter=0,
        textColor=TEXT_MUTED
    )

    def section_header(title):
        return [
            Paragraph(f'<b>{title}</b>', section_title_style),
            HRFlowable(width='100%', thickness=0.8, color=ACCENT,
                       spaceBefore=0, spaceAfter=3),
        ]

    def experience_entry(exp):
        elements = [
            Paragraph(f'<b>{exp["title"]}</b>', job_title_style),
            Paragraph(f'{exp["company"]}  |  {exp["dates"]}  |  {exp["location"]}',
                      job_meta_style),
        ]
        for b in exp["bullets"]:
            elements.append(Paragraph(f'• {b}', bullet_style))
        elements.append(Spacer(1, 2))
        return elements

    doc = SimpleDocTemplate(
        output_path, pagesize=A4,
        leftMargin=1.4*cm, rightMargin=1.4*cm,
        topMargin=1.2*cm, bottomMargin=1.2*cm,
        title='Resume - Roman M. - COO Web3',
        author='Roman M.', creator='Z.ai'
    )

    story = []

    # Header
    story.append(Paragraph(NAME, name_style))
    story.append(Paragraph(TARGET_ROLE, target_role_style))
    story.append(Paragraph(CONTACT_LINE, contact_style))

    # Summary
    story.extend(section_header('PROFESSIONAL SUMMARY'))
    story.append(Paragraph(PROFESSIONAL_SUMMARY, body_style))

    # Skills
    story.extend(section_header('CORE SKILLS'))
    for cat, vals in SKILLS:
        story.append(Paragraph(f'<b>{cat}:</b>  {vals}', skills_style))

    # Experience
    story.extend(section_header('WORK EXPERIENCE'))
    for exp in EXPERIENCE:
        story.extend(experience_entry(exp))

    # Earlier experience (compressed — no bullets)
    if EARLIER_EXPERIENCE:
        story.append(Paragraph('<b>Earlier Experience</b>', job_title_style))
        for title, company, dates, desc in EARLIER_EXPERIENCE:
            story.append(Paragraph(
                f'<b>{title}</b> — {company}  |  {dates}', earlier_style))
            story.append(Paragraph(desc, earlier_style))
        story.append(Spacer(1, 2))

    # Education
    story.extend(section_header('EDUCATION'))
    for degree, school, dates in EDUCATION:
        story.append(Paragraph(f'<b>{degree}</b>', job_title_style))
        story.append(Paragraph(f'{school}  |  {dates}', job_meta_style))

    # Languages
    story.extend(section_header('LANGUAGES'))
    story.append(Paragraph(LANGUAGES, body_style))

    doc.build(story)
    print(f"  [PDF] Generated: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# DOCX GENERATION  (ATS-safe single-column, no tables, no sidebar)
# ─────────────────────────────────────────────────────────────────────────────

def _set_cell_borders_none(cell):
    """Helper to strip borders from a cell — used only when we can't avoid tables."""
    pass  # not used in ATS-safe version


def _add_section_heading(doc, title, accent_color="1A3352"):
    """Add a section heading with bottom border (ATS-safe)."""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(7)
    p.paragraph_format.space_after = Pt(2)
    p.paragraph_format.keep_with_next = True
    run = p.add_run(title.upper())
    run.bold = True
    run.font.size = Pt(12)
    run.font.name = 'Calibri'
    run.font.color.rgb = RGBColor.from_string(accent_color)

    # Bottom border on the paragraph
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')        # 0.75pt
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), accent_color)
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


def _add_body_paragraph(doc, text, font_size=9.5, bold=False, italic=False,
                        color=None, alignment=None, space_after=2, space_before=0):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(space_after)
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.line_spacing = 1.1
    if alignment is not None:
        p.alignment = alignment
    run = p.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    run.bold = bold
    run.italic = italic
    if color is not None:
        run.font.color.rgb = RGBColor.from_string(color)
    return p


def _add_bullet(doc, text, font_size=9.5):
    """ATS-safe bullet: literal '• ' prefix, plain paragraph (no List style)."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0.5)
    p.paragraph_format.first_line_indent = Cm(-0.3)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    run = p.add_run('•  ' + text)
    run.font.name = 'Calibri'
    run.font.size = Pt(font_size)
    return p


def _add_skill_row(doc, category, values):
    """Skill row: bold category label + plain values."""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(0)
    p.paragraph_format.space_after = Pt(0)
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.line_spacing = 1.1
    r1 = p.add_run(f'{category}:  ')
    r1.bold = True
    r1.font.name = 'Calibri'
    r1.font.size = Pt(9.5)
    r1.font.color.rgb = RGBColor.from_string('1A1A1A')
    r2 = p.add_run(values)
    r2.font.name = 'Calibri'
    r2.font.size = Pt(9.5)
    return p


def build_docx(output_path):
    """Generate ATS-optimized DOCX resume (single-column, no tables)."""
    doc = Document()

    # Page setup: A4, tight margins
    section = doc.sections[0]
    section.page_height = Cm(29.7)
    section.page_width = Cm(21.0)
    section.top_margin = Cm(1.2)
    section.bottom_margin = Cm(1.2)
    section.left_margin = Cm(1.4)
    section.right_margin = Cm(1.4)

    # Set default style font
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(9.5)

    # ── Header ──
    # Name
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(NAME)
    run.bold = True
    run.font.name = 'Calibri'
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor.from_string('1A1A1A')

    # Target role
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    run = p.add_run(TARGET_ROLE)
    run.font.name = 'Calibri'
    run.font.size = Pt(10.5)
    run.font.color.rgb = RGBColor.from_string('1A3352')

    # Contact
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(3)
    run = p.add_run(CONTACT_LINE)
    run.font.name = 'Calibri'
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor.from_string('555555')
    run.italic = True

    # ── Professional Summary ──
    _add_section_heading(doc, 'Professional Summary')
    _add_body_paragraph(doc, PROFESSIONAL_SUMMARY, font_size=9.5, space_after=3)

    # ── Core Skills ──
    _add_section_heading(doc, 'Core Skills')
    for cat, vals in SKILLS:
        _add_skill_row(doc, cat, vals)

    # ── Work Experience ──
    _add_section_heading(doc, 'Work Experience')

    for exp in EXPERIENCE:
        # Job title
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(3)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(exp['title'])
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)
        run.font.color.rgb = RGBColor.from_string('1A1A1A')

        # Meta: company | dates | location
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(f'{exp["company"]}  |  {exp["dates"]}  |  {exp["location"]}')
        run.italic = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string('555555')

        # Bullets
        for b in exp['bullets']:
            _add_bullet(doc, b, font_size=9.5)

    # Earlier experience (compressed)
    if EARLIER_EXPERIENCE:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(4)
        p.paragraph_format.keep_with_next = True
        run = p.add_run('Earlier Experience')
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        for title, company, dates, desc in EARLIER_EXPERIENCE:
            p = doc.add_paragraph()
            p.paragraph_format.left_indent = Cm(0.5)
            p.paragraph_format.space_after = Pt(0)
            p.paragraph_format.space_before = Pt(1)
            p.paragraph_format.line_spacing = 1.1
            r1 = p.add_run(f'{title}')
            r1.bold = True
            r1.font.name = 'Calibri'
            r1.font.size = Pt(9)
            r1.font.color.rgb = RGBColor.from_string('555555')
            r2 = p.add_run(f' — {company}  |  {dates}')
            r2.italic = True
            r2.font.name = 'Calibri'
            r2.font.size = Pt(9)
            r2.font.color.rgb = RGBColor.from_string('555555')

            p2 = doc.add_paragraph()
            p2.paragraph_format.left_indent = Cm(0.5)
            p2.paragraph_format.space_after = Pt(1)
            p2.paragraph_format.line_spacing = 1.1
            r = p2.add_run(desc)
            r.font.name = 'Calibri'
            r.font.size = Pt(9)
            r.font.color.rgb = RGBColor.from_string('555555')

    # ── Education ──
    _add_section_heading(doc, 'Education')
    for degree, school, dates in EDUCATION:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(0)
        p.paragraph_format.space_before = Pt(1)
        p.paragraph_format.keep_with_next = True
        run = p.add_run(degree)
        run.bold = True
        run.font.name = 'Calibri'
        run.font.size = Pt(10)

        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(1)
        run = p.add_run(f'{school}  |  {dates}')
        run.italic = True
        run.font.name = 'Calibri'
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor.from_string('555555')

    # ── Languages ──
    _add_section_heading(doc, 'Languages')
    _add_body_paragraph(doc, LANGUAGES, font_size=9.5, space_after=2)

    # Set document metadata
    cp = doc.core_properties
    cp.title = 'Resume - Roman M. - COO Web3'
    cp.author = 'Roman M.'
    cp.subject = 'Resume'
    cp.comments = 'ATS-optimized resume targeting COO / Head of Operations roles at Series A Web3 startups'

    doc.save(output_path)
    print(f"  [DOCX] Generated: {output_path}")


# ─────────────────────────────────────────────────────────────────────────────
# MAIN
# ─────────────────────────────────────────────────────────────────────────────

if __name__ == '__main__':
    out_dir = '/home/z/my-project/download'
    os.makedirs(out_dir, exist_ok=True)

    pdf_path = os.path.join(out_dir, 'Roman_M_Resume_COO_Web3.pdf')
    docx_path = os.path.join(out_dir, 'Roman_M_Resume_COO_Web3.docx')

    print("Generating ATS-optimized resume for Roman M....")
    build_pdf(pdf_path)
    build_docx(docx_path)
    print("\nDone. Deliverables:")
    print(f"  - {pdf_path}")
    print(f"  - {docx_path}")
